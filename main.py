import streamlit as st
import tempfile
import base64
from quiz_engine import generate_questions_from_text
from doc_processor import read_docx, convert_to_docx
from docx import Document


st.set_page_config(page_title="Quiz Generator", layout="wide")
st.markdown("""
<link href="https://fonts.googleapis.com/css2?family=Pacifico&display=swap" rel="stylesheet">
<h1 style='text-align: center; font-family: "Pacifico", cursive; font-size: 80px; color: #2c3e50;'>TheQuizLab</h1>
""", unsafe_allow_html=True)

with st.sidebar:
    st.header("Settings")
    difficulty = st.selectbox("Difficulty Level", ["Easy", "Medium", "Hard"])
    question_types = st.multiselect("Question Types", ["MCQ", "Fill-in-the-blank"], default=["MCQ"])
    num_questions = st.number_input("Number of Questions", min_value=1, max_value=100, value=3, step=1)

with st.expander("💡 Tips for better results"):
    st.markdown("""
    - Upload documents with clear, well-structured text  
    - For MCQ questions, documents with factual info work best  
    - Use medium-length documents (1–5 pages)  
    """)

uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])
conversion_file = st.file_uploader("Not a DOCX? Upload your file here", type=["pdf", "txt", "csv", "ipynb"])

def generate_quiz(text, num_questions, difficulty, question_type):
    max_retries = 3
    for attempt in range(max_retries):
        quiz = generate_questions_from_text(text, num_questions, difficulty, question_type)
        if quiz:
            return quiz
        st.warning(f"Attempt {attempt + 1} failed. Retrying...")
    return None

def get_docx_download_button(quiz_data, filename="quiz_output.docx"):
    doc = Document()
    doc.add_heading("Generated Quiz", level=1)

    for i, q in enumerate(quiz_data):
        doc.add_paragraph(f"Q{i+1}: {q.get('question', 'No question provided')}", style='List Number')
        if q.get("type") == "MCQ" and "options" in q:
            for j, opt in enumerate(q["options"]):
                doc.add_paragraph(f"{chr(97+j)}. {opt}", style='List Bullet')
        doc.add_paragraph(f"Answer: {q.get('answer', 'No answer provided')}", style='Intense Quote')
        doc.add_paragraph("")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp.seek(0)
        b64 = base64.b64encode(tmp.read()).decode()
        href = f'''
            <a href="data:application/octet-stream;base64,{b64}" 
               download="{filename}" 
               style="text-decoration: none; font-size: 18px;">
               <button style='padding:8px 16px;border:none;background-color:#3498db;color:white;border-radius:6px;cursor:pointer;'>
               Download DOCX
               </button>
            </a>
        '''
        return href

if st.button("Generate Quiz") and (uploaded_file or conversion_file):
    text = ""
    if uploaded_file:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(uploaded_file.read())
            text = read_docx(tmp.name)
    elif conversion_file:
        file_ext = conversion_file.name.split(".")[-1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp:
            tmp.write(conversion_file.read())
            tmp_path = tmp.name
        try:
            text = read_docx(convert_to_docx(tmp_path))
        except Exception as e:
            st.error(f"File conversion failed: {str(e)}")
            st.stop()

    if not text or len(text.split()) < 50:
        st.error("Document is too short or empty. Please upload a longer document (min 50 words).")
        st.stop()

    with st.spinner("Generating questions..."):
        progress_bar = st.progress(0)
        quiz = generate_quiz(text, num_questions, difficulty, question_types[0] if question_types else "MCQ")
        progress_bar.progress(100)

        if quiz:
            st.success(f"Generated {len(quiz)} questions!")
            for i, q in enumerate(quiz):
                st.subheader(f"Q{i+1}: {q.get('question', 'No question provided')}")

                if q.get("type") == "MCQ" and "options" in q:
                    cols = st.columns(2)
                    for j, opt in enumerate(q["options"]):
                        cols[j % 2].write(f"{chr(97+j)}. {opt}")
                with st.expander("Show Answer"):
                    st.markdown(f"Answer: {q.get('answer', 'No answer provided')}")
                st.divider()

            cols = st.columns([4, 1]) 
            with cols[1]:
                st.markdown(get_docx_download_button(quiz), unsafe_allow_html=True)

        else:
            st.error("Failed to generate valid questions. Please try again.")