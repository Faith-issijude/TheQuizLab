from docx import Document
import os
import pypandoc
import tempfile
from pdfminer.high_level import extract_text

def read_docx(file_path: str) -> str:
    """Extracts text from .docx files."""
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def pdf_to_docx(file_path: str) -> str:
    """
    Converts PDF to DOCX using python-docx and pdfminer.six
    (Pure Python solution - no external dependencies)
    """
    try:
        text = extract_text(file_path)
        output_path = tempfile.mktemp(suffix=".docx")
        doc = Document()
        
        # Preserve some formatting by splitting paragraphs
        for paragraph in text.split('\n\n'):
            if paragraph.strip():  # Skip empty paragraphs
                doc.add_paragraph(paragraph.strip())
        
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise ValueError(f"PDF to DOCX conversion failed: {str(e)}")

def convert_to_docx(file_path: str) -> str:
    """
    Universal converter that handles:
    - PDF files (using pure Python)
    - Other supported formats (using pandoc)
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Handle PDF with pure Python solution
    if file_ext == '.pdf':
        return pdf_to_docx(file_path)
    
    # Handle other formats with pandoc
    output_path = tempfile.mktemp(suffix=".docx")
    try:
        pypandoc.convert_file(file_path, 'docx', outputfile=output_path)
        return output_path
    except RuntimeError as e:
        if "Invalid input format" in str(e):
            supported = pypandoc.get_pandoc_formats()[0]
            raise ValueError(
                f"Unsupported file format '{file_ext}'. "
                f"Supported input formats: {', '.join(supported)}"
            )
        raise ValueError(f"Conversion failed: {str(e)}")

def get_supported_formats() -> tuple[list[str], list[str]]:
    """Returns tuple of (input_formats, output_formats)"""
    try:
        return pypandoc.get_pandoc_formats()
    except:
        # Fallback if pandoc isn't available
        return (['pdf', 'docx'], ['docx'])